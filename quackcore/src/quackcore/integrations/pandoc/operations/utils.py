# quackcore/src/quackcore/integrations/pandoc/operations/utils.py
"""
Utility functions for pandoc _operations.

This module provides helper functions for pandoc conversion _operations,
such as validation, metrics tracking, and pandoc installation verification.
All file path values are handled as strings. Filesystem _operations are delegated
to the quackcore.fs service.
"""

import sys
import time
import types
from typing import Any

from quackcore.errors import QuackIntegrationError
from quackcore.integrations.pandoc.config import PandocConfig
from quackcore.integrations.pandoc.models import ConversionMetrics, FileInfo
from quackcore.logging import get_logger

logger = get_logger(__name__)

# Ensure fs module is properly available
if 'quackcore.fs.service' not in sys.modules:
    # Create the module hierarchy if needed
    if 'quackcore' not in sys.modules:
        quackcore_mod = types.ModuleType('quackcore')
        sys.modules['quackcore'] = quackcore_mod

    if 'quackcore.fs' not in sys.modules:
        fs_mod = types.ModuleType('quackcore.fs')
        sys.modules['quackcore.fs'] = fs_mod

    service_mod = types.ModuleType('quackcore.fs.service')
    service_mod.standalone = types.SimpleNamespace()
    sys.modules['quackcore.fs.service'] = service_mod

from quackcore.fs.service import standalone as fs


def verify_pandoc() -> str:
    """
    Verify pandoc installation and version.

    Returns:
        str: Pandoc version string.

    Raises:
        QuackIntegrationError: If pandoc is not installed.
    """
    try:
        import importlib
        pypandoc = importlib.import_module('pypandoc')
        version = pypandoc.get_pandoc_version()
        logger.info(f"Found pandoc version: {version}")
        return version
    except ImportError as err:
        msg = "pypandoc module is not installed"
        logger.error(msg)
        raise QuackIntegrationError(msg, {"module": "pypandoc"}) from err
    except OSError as err:
        msg = "Pandoc is not installed. Please install pandoc first."
        logger.error(msg)
        raise QuackIntegrationError(msg, {"original_error": str(err)}) from err
    except Exception as e:
        msg = f"Error checking pandoc: {str(e)}"
        logger.error(msg)
        raise QuackIntegrationError(msg, {"original_error": str(e)}) from e


def prepare_pandoc_args(
        config: PandocConfig,
        source_format: str,
        target_format: str,
        extra_args: list[str] | None = None,
) -> list[str]:
    """
    Prepare pandoc conversion arguments.

    Args:
        config: Conversion configuration.
        source_format: Source format.
        target_format: Target format.
        extra_args: Additional arguments.

    Returns:
        list[str]: List of pandoc arguments.
    """
    pandoc_opts = config.pandoc_options
    raw_args = [
        f"--wrap={pandoc_opts.wrap}",
        "--standalone" if pandoc_opts.standalone else None,
        f"--markdown-headings={pandoc_opts.markdown_headings}",
        "--reference-links" if pandoc_opts.reference_links else None,
    ]
    args: list[str] = [arg for arg in raw_args if arg is not None]

    # Convert resource paths to strings.
    for res_path in pandoc_opts.resource_path:
        args.append(f"--resource-path={str(res_path)}")

    if source_format == "html" and target_format == "markdown":
        args.extend(config.html_to_md_extra_args)
    elif source_format == "markdown" and target_format == "docx":
        args.extend(config.md_to_docx_extra_args)
    if extra_args:
        args.extend(extra_args)

    return args


def validate_html_structure(
        content: str, check_links: bool = False
) -> tuple[bool, list[str]]:
    """
    Validate HTML document structure.

    Args:
        content: HTML content.
        check_links: Whether to check links.

    Returns:
        tuple: (is_valid, list of error messages).
    """
    errors: list[str] = []
    try:
        # Attempt to import BeautifulSoup
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            errors.append("HTML validation error: bs4 module not installed")
            return False, errors

        soup = BeautifulSoup(content, "html.parser")
        if not soup.find("body"):
            errors.append("HTML document missing body tag")
            return False, errors

        if not (
                soup.find(["h1", "h2", "h3", "h4", "h5", "h6"])
                or soup.find(["header", "section", "article"])
        ):
            logger.warning("HTML document has no header tags or structural elements")

        if check_links:
            links = soup.find_all("a")
            empty_links = [
                str(link) for link in links if not (link.get("href") or "").strip()
            ]
            if empty_links:
                errors.append(f"Found {len(empty_links)} empty links in document")
        return len(errors) == 0, errors
    except Exception as e:
        errors.append(f"HTML validation error: {str(e)}")
        return False, errors


def validate_docx_structure(
        docx_path: str, check_links: bool = False
) -> tuple[bool, list[str]]:
    """
    Validate DOCX document structure.

    Args:
        docx_path: Path to DOCX file (as a string).
        check_links: Whether to check links.

    Returns:
        tuple: (is_valid, list of error messages).
    """
    errors: list[str] = []
    try:
        # Attempt to import docx module
        try:
            import docx as docx_module
        except ImportError:
            logger.warning("python-docx module is not installed")
            return True, []  # Return valid if docx module not available

        doc = docx_module.Document(docx_path)
        if len(doc.paragraphs) == 0:
            errors.append("DOCX document has no paragraphs")
            return False, errors

        has_heading = any(
            para.style and para.style.name.startswith("Heading")
            for para in doc.paragraphs
        )
        if not has_heading:
            logger.warning("DOCX document has no heading styles")

        if check_links:
            if not hasattr(doc, "part") or doc.part is None:
                errors.append("Document structure appears incomplete")
                return False, errors

        return len(errors) == 0, errors
    except ImportError:
        logger.warning("python-docx module is not installed")
        return True, []  # Return valid if docx module not available
    except Exception as e:
        errors.append(f"DOCX validation error: {str(e)}")
        return False, errors


def safe_convert_to_int(value: Any, default: int = 0) -> int:
    """
    Safely convert a value to an integer with fallback to default.

    Args:
        value: The value to convert.
        default: Default value if conversion fails.

    Returns:
        int: Converted integer or default value.
    """
    if value is None:
        return default

    try:
        return int(value)
    except (ValueError, TypeError):
        logger.warning(
            f"Could not convert value to integer: {value}, using default {default}")
        return default


def get_size_str_wrapper(size: int) -> str:
    """
    Wrapper for fs.get_file_size_str that handles errors and returns a string.

    Args:
        size: Size in bytes

    Returns:
        str: Human-readable size string
    """
    try:
        result = fs.get_file_size_str(size)
        if hasattr(result, 'data') and result.success:
            return result.data
        return f"{size}B"
    except Exception as e:
        logger.warning(f"Error getting file size string: {e}")
        return f"{size}B"


def check_file_size(
        converted_size: int | None, validation_min_size: int | None
) -> tuple[bool, list[str]]:
    """
    Check if the converted file meets the minimum file size.

    Args:
        converted_size: Size of the converted file.
        validation_min_size: Minimum file size threshold.

    Returns:
        tuple: (is_valid, list of error messages).
    """
    errors: list[str] = []
    converted_size_int = safe_convert_to_int(converted_size, 0)
    validation_min_size_int = safe_convert_to_int(validation_min_size, 0)

    if validation_min_size_int > 0 and converted_size_int < validation_min_size_int:
        converted_size_str = get_size_str_wrapper(converted_size_int)
        min_size_str = get_size_str_wrapper(validation_min_size_int)
        errors.append(
            f"Converted file size ({converted_size_str}) is below the minimum threshold ({min_size_str})"
        )
        return False, errors
    return True, errors


def check_conversion_ratio(
        converted_size: int | None, original_size: int | None, threshold: float | None
) -> tuple[bool, list[str]]:
    """
    Check if the converted file size is not drastically smaller than the original.

    Args:
        converted_size: Size of the converted file.
        original_size: Size of the original file.
        threshold: Minimum ratio threshold.

    Returns:
        tuple: (is_valid, list of error messages).
    """
    errors: list[str] = []
    converted_size_int = safe_convert_to_int(converted_size, 0)
    original_size_int = safe_convert_to_int(original_size, 0)
    threshold_float = float(threshold) if threshold is not None else 0.1

    if original_size_int > 0:
        conversion_ratio = converted_size_int / original_size_int
        if conversion_ratio < threshold_float:
            converted_size_str = get_size_str_wrapper(converted_size_int)
            original_size_str = get_size_str_wrapper(original_size_int)
            errors.append(
                f"Conversion error: Converted file size ({converted_size_str}) is less than "
                f"{threshold_float * 100:.0f}% of the original file size ({original_size_str}) (ratio: {conversion_ratio:.2f})."
            )
            return False, errors
    return True, errors


def track_metrics(
        filename: str,
        start_time: float,
        original_size: int,
        converted_size: int,
        metrics: ConversionMetrics,
        config: PandocConfig,
) -> None:
    """
    Track conversion metrics.

    Args:
        filename: Name of the file (as a string).
        start_time: Start time of conversion.
        original_size: Size of the original file.
        converted_size: Size of the converted file.
        metrics: Metrics tracker.
        config: Configuration object.
    """
    if config.metrics.track_conversion_time:
        end_time = time.time()
        duration = end_time - start_time
        metrics.conversion_times[filename] = {"start": start_time, "end": end_time}
        logger.info(f"Conversion time for {filename}: {duration:.2f} seconds")

    if config.metrics.track_file_sizes:
        original_size_int = safe_convert_to_int(original_size, 0)
        converted_size_int = safe_convert_to_int(converted_size, 0)
        ratio = converted_size_int / original_size_int if original_size_int > 0 else 0
        metrics.file_sizes[filename] = {
            "original": original_size_int,
            "converted": converted_size_int,
            "ratio": ratio,
        }

        # Handle file size strings safely
        original_size_str = get_size_str_wrapper(original_size_int)
        converted_size_str = get_size_str_wrapper(converted_size_int)

        logger.info(
            f"File size change for {filename}: {original_size_str} -> {converted_size_str}"
        )


def get_file_info(path: str, format_hint: str | None = None) -> FileInfo:
    """
    Get file information for conversion.

    Args:
        path: Path to the file (as a string).
        format_hint: Hint about the file format.

    Returns:
        FileInfo: File information.

    Raises:
        QuackIntegrationError: If the file does not exist.
    """
    file_info = fs.get_file_info(path)
    if not file_info.success or not file_info.exists:
        raise QuackIntegrationError(f"File not found: {path}")

    # Convert file size to integer safely
    file_size = safe_convert_to_int(file_info.size, 0)
    modified_time: float | None = file_info.modified

    # Determine format name
    if format_hint:
        format_name = format_hint
    else:
        try:
            # Get extension safely
            ext_result = fs.get_extension(path)
            extension = ""

            if hasattr(ext_result, 'success') and ext_result.success and hasattr(
                    ext_result, 'data'):
                extension = ext_result.data
            else:
                # Fallback extension extraction if get_extension not working correctly
                extension = path.split('.')[-1] if isinstance(path,
                                                              str) and '.' in path else ""

            # Map the extension to a format name
            mapping: dict[str, str] = {
                "md": "markdown",
                "markdown": "markdown",
                "html": "html",
                "htm": "html",
                "docx": "docx",
                "doc": "docx",
                "pdf": "pdf",
                "txt": "plain",
            }
            format_name = mapping.get(extension, extension)
        except Exception as e:
            logger.warning(f"Error getting file extension: {e}. Using fallback.")
            # Fallback to extension from path
            extension = path.split('.')[-1] if isinstance(path,
                                                          str) and '.' in path else "unknown"
            format_name = extension

    return FileInfo(
        path=path,
        format=format_name,
        size=file_size,
        modified=modified_time,
        extra_args=[],
    )
