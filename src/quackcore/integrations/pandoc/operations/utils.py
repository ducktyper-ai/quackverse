# src/quackcore/integrations/pandoc/_operations/utils.py
"""
Utility functions for pandoc _operations.

This module provides helper functions for pandoc conversion _operations,
such as validation, metrics tracking, and pandoc installation verification.
All file path values are handled as strings. Filesystem _operations are delegated
to the quackcore.fs service.
"""

import time

from quackcore.errors import QuackIntegrationError
from quackcore.fs import service as fs  # fs functions work with string paths
from quackcore.integrations.pandoc.config import PandocConfig
from quackcore.integrations.pandoc.models import ConversionMetrics, FileInfo
from quackcore.logging import get_logger

logger = get_logger(__name__)


def verify_pandoc() -> str:
    """
    Verify pandoc installation and version.

    Returns:
        str: Pandoc version string.

    Raises:
        QuackIntegrationError: If pandoc is not installed.
    """
    try:
        import pypandoc

        version = pypandoc.get_pandoc_version()
        logger.info(f"Found pandoc version: {version}")
        return version
    except ImportError as err:
        msg = "pypandoc module is not installed"
        logger.error(msg)
        raise QuackIntegrationError(msg, {"module": "pypandoc"}) from err
    except OSError as err:
        msg = "Pandoc is not installed. Please install pandoc first."
        logger.error(msg)
        raise QuackIntegrationError(msg, {"original_error": str(err)}) from err
    except Exception as e:
        msg = f"Error checking pandoc: {str(e)}"
        logger.error(msg)
        raise QuackIntegrationError(msg, {"original_error": str(e)}) from e


def prepare_pandoc_args(
    config: PandocConfig,
    source_format: str,
    target_format: str,
    extra_args: list[str] | None = None,
) -> list[str]:
    """
    Prepare pandoc conversion arguments.

    Args:
        config: Conversion configuration.
        source_format: Source format.
        target_format: Target format.
        extra_args: Additional arguments.

    Returns:
        list[str]: List of pandoc arguments.
    """
    pandoc_opts = config.pandoc_options
    raw_args = [
        f"--wrap={pandoc_opts.wrap}",
        "--standalone" if pandoc_opts.standalone else None,
        f"--markdown-headings={pandoc_opts.markdown_headings}",
        "--reference-links" if pandoc_opts.reference_links else None,
    ]
    args: list[str] = [arg for arg in raw_args if arg is not None]

    # Convert resource paths to strings.
    for res_path in pandoc_opts.resource_path:
        args.append(f"--resource-path={str(res_path)}")

    if source_format == "html" and target_format == "markdown":
        args.extend(config.html_to_md_extra_args)
    elif source_format == "markdown" and target_format == "docx":
        args.extend(config.md_to_docx_extra_args)
    if extra_args:
        args.extend(extra_args)

    return args


def validate_html_structure(
    content: str, check_links: bool = False
) -> tuple[bool, list[str]]:
    """
    Validate HTML document structure.

    Args:
        content: HTML content.
        check_links: Whether to check links.

    Returns:
        tuple: (is_valid, list of error messages).
    """
    errors: list[str] = []
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(content, "html.parser")
        if not soup.find("body"):
            errors.append("HTML document missing body tag")
            return False, errors

        if not (
            soup.find(["h1", "h2", "h3", "h4", "h5", "h6"])
            or soup.find(["header", "section", "article"])
        ):
            logger.warning("HTML document has no header tags or structural elements")

        if check_links:
            links = soup.find_all("a")
            empty_links = [
                str(link) for link in links if not (link.get("href") or "").strip()
            ]
            if empty_links:
                errors.append(f"Found {len(empty_links)} empty links in document")
        return len(errors) == 0, errors
    except Exception as e:
        errors.append(f"HTML validation error: {str(e)}")
        return False, errors


def validate_docx_structure(
    docx_path: str, check_links: bool = False
) -> tuple[bool, list[str]]:
    """
    Validate DOCX document structure.

    Args:
        docx_path: Path to DOCX file (as a string).
        check_links: Whether to check links.

    Returns:
        tuple: (is_valid, list of error messages).
    """
    errors: list[str] = []
    try:
        from docx import Document

        doc = Document(docx_path)
        if len(doc.paragraphs) == 0:
            errors.append("DOCX document has no paragraphs")
            return False, errors

        has_heading = any(
            para.style and para.style.name.startswith("Heading")
            for para in doc.paragraphs
        )
        if not has_heading:
            logger.warning("DOCX document has no heading styles")

        if check_links:
            if not hasattr(doc, "part") or doc.part is None:
                errors.append("Document structure appears incomplete")
                return False, errors

        return len(errors) == 0, errors
    except ImportError:
        logger.warning("python-docx module is not installed")
        return True, []
    except Exception as e:
        errors.append(f"DOCX validation error: {str(e)}")
        return False, errors


def track_metrics(
    filename: str,
    start_time: float,
    original_size: int,
    converted_size: int,
    metrics: ConversionMetrics,
    config: PandocConfig,
) -> None:
    """
    Track conversion metrics.

    Args:
        filename: Name of the file (as a string).
        start_time: Start time of conversion.
        original_size: Size of the original file.
        converted_size: Size of the converted file.
        metrics: Metrics tracker.
        config: Configuration object.
    """
    if config.metrics.track_conversion_time:
        end_time = time.time()
        duration = end_time - start_time
        metrics.conversion_times[filename] = {"start": start_time, "end": end_time}
        logger.info(f"Conversion time for {filename}: {duration:.2f} seconds")

    if config.metrics.track_file_sizes:
        original_size_int = int(original_size) if original_size is not None else 0
        converted_size_int = int(converted_size) if converted_size is not None else 0
        ratio = converted_size_int / original_size_int if original_size_int > 0 else 0
        metrics.file_sizes[filename] = {
            "original": original_size_int,
            "converted": converted_size_int,
            "ratio": ratio,
        }
        original_size_str = fs.get_file_size_str(original_size_int)
        converted_size_str = fs.get_file_size_str(converted_size_int)
        logger.info(
            f"File size change for {filename}: {original_size_str} -> {converted_size_str}"
        )


def get_file_info(path: str, format_hint: str | None = None) -> FileInfo:
    """
    Get file information for conversion.

    Args:
        path: Path to the file (as a string).
        format_hint: Hint about the file format.

    Returns:
        FileInfo: File information.

    Raises:
        QuackIntegrationError: If the file does not exist.
    """
    file_info = fs.get_file_info(path)
    if not file_info.success or not file_info.exists:
        raise QuackIntegrationError(f"File not found: {path}")
    try:
        file_size: int = int(file_info.size) if file_info.size is not None else 0
    except (TypeError, ValueError):
        file_size = 1024
        logger.warning(
            f"Could not convert file size to integer: {file_info.size}, using default"
        )
    modified_time: float | None = file_info.modified
    if format_hint:
        format_name = format_hint
    else:
        extension = fs.get_extension(path)
        mapping: dict[str, str] = {
            "md": "markdown",
            "markdown": "markdown",
            "html": "html",
            "htm": "html",
            "docx": "docx",
            "doc": "docx",
            "pdf": "pdf",
            "txt": "plain",
        }
        format_name = mapping.get(extension, extension)
    return FileInfo(
        path=path,
        format=format_name,
        size=file_size,
        modified=modified_time,
        extra_args=[],
    )


def check_file_size(
    converted_size: int, validation_min_size: int
) -> tuple[bool, list[str]]:
    """
    Check if the converted file meets the minimum file size.

    Args:
        converted_size: Size of the converted file.
        validation_min_size: Minimum file size threshold.

    Returns:
        tuple: (is_valid, list of error messages).
    """
    errors: list[str] = []
    converted_size_int = int(converted_size) if converted_size is not None else 0
    validation_min_size_int = (
        int(validation_min_size) if validation_min_size is not None else 0
    )
    if validation_min_size_int > 0 and converted_size_int < validation_min_size_int:
        converted_size_str = fs.get_file_size_str(converted_size_int)
        min_size_str = fs.get_file_size_str(validation_min_size_int)
        errors.append(
            f"Converted file size ({converted_size_str}) is below the minimum threshold ({min_size_str})"
        )
        return False, errors
    return True, errors


def check_conversion_ratio(
    converted_size: int, original_size: int, threshold: float
) -> tuple[bool, list[str]]:
    """
    Check if the converted file size is not drastically smaller than the original.

    Args:
        converted_size: Size of the converted file.
        original_size: Size of the original file.
        threshold: Minimum ratio threshold.

    Returns:
        tuple: (is_valid, list of error messages).
    """
    errors: list[str] = []
    converted_size_int = int(converted_size) if converted_size is not None else 0
    original_size_int = int(original_size) if original_size is not None else 0
    threshold_float = float(threshold) if threshold is not None else 0.1
    if original_size_int > 0:
        conversion_ratio = converted_size_int / original_size_int
        if conversion_ratio < threshold_float:
            converted_size_str = fs.get_file_size_str(converted_size_int)
            original_size_str = fs.get_file_size_str(original_size_int)
            errors.append(
                f"Conversion error: Converted file size ({converted_size_str}) is less than "
                f"{threshold_float * 100:.0f}% of the original file size ({original_size_str}) (ratio: {conversion_ratio:.2f})."
            )
            return False, errors
    return True, errors
